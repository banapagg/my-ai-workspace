#!/usr/bin/env python3
"""
Document Writer Worker - ドキュメント作成ワーカー

役割: Word/PDF文書の生成、体裁の良いフォーマット適用
対象: コンサルタント、営業職、事務職
"""

import json
import os
from datetime import datetime
from typing import Dict, List, Any, Optional

# python-docxが利用可能かチェック
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("[!] python-docx がインストールされていません")
    print("    インストール: pip install python-docx")


class DocumentWriterWorker:
    """ドキュメント作成ワーカー"""

    def __init__(self):
        """初期化"""
        self.name = "document_writer"
        self.description = "Word/PDF文書の作成を担当"
        self.skills = ["document-formatting", "business-writing"]
        self.templates = self._load_templates()

    def _load_templates(self) -> Dict[str, Any]:
        """テンプレート定義をロード"""
        return {
            "proposal": {
                "name": "提案書・企画書",
                "sections": [
                    {"title": "表紙", "type": "cover"},
                    {"title": "目次", "type": "toc"},
                    {"title": "背景・課題", "type": "content"},
                    {"title": "提案内容", "type": "content"},
                    {"title": "実施計画", "type": "content"},
                    {"title": "期待効果", "type": "content"},
                    {"title": "費用見積", "type": "table"},
                    {"title": "付録", "type": "appendix"}
                ]
            },
            "report": {
                "name": "報告書・レポート",
                "sections": [
                    {"title": "表紙", "type": "cover"},
                    {"title": "目次", "type": "toc"},
                    {"title": "概要", "type": "content"},
                    {"title": "調査結果", "type": "content"},
                    {"title": "分析", "type": "content"},
                    {"title": "結論・提言", "type": "content"},
                    {"title": "付録", "type": "appendix"}
                ]
            },
            "minutes": {
                "name": "議事録",
                "sections": [
                    {"title": "会議情報", "type": "header"},
                    {"title": "出席者", "type": "list"},
                    {"title": "議題", "type": "content"},
                    {"title": "決定事項", "type": "list"},
                    {"title": "アクション項目", "type": "table"},
                    {"title": "次回予定", "type": "footer"}
                ]
            }
        }

    def execute(self, task: Dict[str, Any]) -> Dict[str, Any]:
        """タスクを実行"""
        print(f"\n[*] {self.name}: タスクを実行中...")
        print(f"   タスク: {task.get('description', 'N/A')}")

        task_type = task.get("type", "create_document")

        if task_type == "create_document":
            return self._create_document(task)
        elif task_type == "suggest_structure":
            return self._suggest_structure(task)
        elif task_type == "format_content":
            return self._format_content(task)
        else:
            return {
                "status": "error",
                "worker": self.name,
                "error": f"不明なタスクタイプ: {task_type}"
            }

    def _suggest_structure(self, task: Dict[str, Any]) -> Dict[str, Any]:
        """文書構成を提案"""
        doc_type = task.get("doc_type", "report")
        topic = task.get("topic", "")
        purpose = task.get("purpose", "")

        template = self.templates.get(doc_type, self.templates["report"])

        structure = {
            "doc_type": doc_type,
            "template_name": template["name"],
            "topic": topic,
            "purpose": purpose,
            "sections": template["sections"],
            "recommendations": self._get_recommendations(doc_type)
        }

        result = {
            "status": "success",
            "worker": self.name,
            "output": structure,
            "logs": [
                f"文書タイプ: {template['name']}",
                f"セクション数: {len(template['sections'])}",
                "構成を提案しました"
            ]
        }

        print(f"   [+] 文書構成を提案しました（{len(template['sections'])}セクション）")
        return result

    def _get_recommendations(self, doc_type: str) -> List[str]:
        """文書タイプ別の推奨事項"""
        recommendations = {
            "proposal": [
                "表紙には企業ロゴと日付を記載",
                "目次は自動生成機能を使用",
                "費用見積は表形式で明確に",
                "実施スケジュールはガントチャート形式で"
            ],
            "report": [
                "概要（エグゼクティブサマリー）は1ページ以内",
                "データは図表で視覚化",
                "結論は箇条書きで明確に",
                "出典を必ず記載"
            ],
            "minutes": [
                "会議開始・終了時刻を記録",
                "決定事項と保留事項を明確に区別",
                "アクション項目には担当者と期限を記載",
                "会議後24時間以内に配布"
            ]
        }

        return recommendations.get(doc_type, [])

    def _create_document(self, task: Dict[str, Any]) -> Dict[str, Any]:
        """Word文書を作成"""
        if not DOCX_AVAILABLE:
            return {
                "status": "error",
                "worker": self.name,
                "error": "python-docx がインストールされていません",
                "suggestion": "pip install python-docx を実行してください"
            }

        doc_type = task.get("doc_type", "report")
        title = task.get("title", "文書")
        content = task.get("content", {})
        output_path = task.get("output_path", f"{title}.docx")

        try:
            doc = Document()

            # ドキュメント設定
            self._setup_document_styles(doc)

            # 表紙
            self._add_cover_page(doc, title, content.get("metadata", {}))

            # 改ページ
            doc.add_page_break()

            # 目次（プレースホルダー）
            self._add_toc_placeholder(doc)

            # 改ページ
            doc.add_page_break()

            # コンテンツセクション
            sections = content.get("sections", [])
            for section in sections:
                self._add_section(doc, section)

            # 保存
            doc.save(output_path)

            result = {
                "status": "success",
                "worker": self.name,
                "output": {
                    "file_path": output_path,
                    "doc_type": doc_type,
                    "sections": len(sections),
                    "file_size": os.path.getsize(output_path) if os.path.exists(output_path) else 0
                },
                "logs": [
                    f"Word文書を作成: {output_path}",
                    f"セクション数: {len(sections)}"
                ]
            }

            print(f"   [+] Word文書作成完了: {output_path}")
            return result

        except Exception as e:
            return {
                "status": "error",
                "worker": self.name,
                "error": str(e)
            }

    def _setup_document_styles(self, doc: Any):
        """ドキュメントのスタイルを設定"""
        # 既存のスタイルを使用
        pass

    def _add_cover_page(self, doc: Any, title: str, metadata: Dict[str, Any]):
        """表紙を追加"""
        # タイトル
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.font.size = Pt(28)
        title_run.font.bold = True
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 余白
        for _ in range(5):
            doc.add_paragraph()

        # メタデータ
        if "company" in metadata:
            company_para = doc.add_paragraph()
            company_run = company_para.add_run(metadata["company"])
            company_run.font.size = Pt(16)
            company_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if "date" in metadata:
            date_para = doc.add_paragraph()
            date_run = date_para.add_run(metadata["date"])
            date_run.font.size = Pt(14)
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            date_para = doc.add_paragraph()
            date_run = date_para.add_run(datetime.now().strftime("%Y年%m月%d日"))
            date_run.font.size = Pt(14)
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if "author" in metadata:
            author_para = doc.add_paragraph()
            author_run = author_para.add_run(f"作成者: {metadata['author']}")
            author_run.font.size = Pt(12)
            author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _add_toc_placeholder(self, doc: Any):
        """目次のプレースホルダーを追加"""
        toc_heading = doc.add_heading("目次", level=1)

        note = doc.add_paragraph()
        note_run = note.add_run("[Wordで目次を更新してください: 参考資料 > 目次の更新]")
        note_run.font.size = Pt(10)
        note_run.font.color.rgb = RGBColor(128, 128, 128)

        # 実際の目次はWordの機能で生成する必要がある
        # python-docxでは目次フィールドの挿入に制限がある

    def _add_section(self, doc: Any, section: Dict[str, Any]):
        """セクションを追加"""
        section_type = section.get("type", "content")
        title = section.get("title", "")
        content = section.get("content", "")

        # 見出し
        if title:
            doc.add_heading(title, level=1)

        # コンテンツタイプに応じた処理
        if section_type == "content":
            # テキストコンテンツ
            if isinstance(content, str):
                paragraphs = content.split("\n")
                for para in paragraphs:
                    if para.strip():
                        doc.add_paragraph(para.strip())

            elif isinstance(content, list):
                for item in content:
                    doc.add_paragraph(item, style='List Bullet')

        elif section_type == "table":
            # 表を追加
            table_data = section.get("table_data", [])
            if table_data:
                self._add_table(doc, table_data)

        elif section_type == "list":
            # 箇条書き
            items = content if isinstance(content, list) else [content]
            for item in items:
                doc.add_paragraph(item, style='List Bullet')

        # セクション間に余白
        doc.add_paragraph()

    def _add_table(self, doc: Any, table_data: List[List[str]]):
        """表を追加"""
        if not table_data:
            return

        rows = len(table_data)
        cols = len(table_data[0]) if rows > 0 else 0

        table = doc.add_table(rows=rows, cols=cols)
        table.style = 'Light Grid Accent 1'

        # データを入力
        for i, row_data in enumerate(table_data):
            row = table.rows[i]
            for j, cell_data in enumerate(row_data):
                row.cells[j].text = str(cell_data)

                # ヘッダー行は太字
                if i == 0:
                    row.cells[j].paragraphs[0].runs[0].font.bold = True

    def _format_content(self, task: Dict[str, Any]) -> Dict[str, Any]:
        """コンテンツのフォーマット提案"""
        content = task.get("content", "")
        doc_type = task.get("doc_type", "report")

        # シンプルなフォーマット提案
        formatted = {
            "original": content,
            "suggestions": [
                "見出しレベルを適切に設定",
                "段落間に適切な余白を挿入",
                "重要なポイントは太字で強調",
                "数値データは表形式で整理"
            ],
            "formatted": content  # 実際のフォーマット処理は複雑なため省略
        }

        result = {
            "status": "success",
            "worker": self.name,
            "output": formatted,
            "logs": ["コンテンツのフォーマット提案を生成"]
        }

        return result


def main():
    """テスト用メイン関数"""
    worker = DocumentWriterWorker()

    # テスト1: 文書構成の提案
    print("\n" + "="*60)
    print("テスト1: 文書構成の提案（報告書）")
    print("="*60)

    task1 = {
        "type": "suggest_structure",
        "doc_type": "report",
        "topic": "市場調査報告",
        "purpose": "経営陣への報告"
    }

    result1 = worker.execute(task1)

    if result1["status"] == "success":
        print(f"\n文書タイプ: {result1['output']['template_name']}")
        print("\nセクション構成:")
        for i, section in enumerate(result1['output']['sections'], 1):
            print(f"  {i}. {section['title']} ({section['type']})")

        print("\n推奨事項:")
        for rec in result1['output']['recommendations']:
            print(f"  - {rec}")

    # テスト2: 議事録の構成提案
    print("\n" + "="*60)
    print("テスト2: 文書構成の提案（議事録）")
    print("="*60)

    task2 = {
        "type": "suggest_structure",
        "doc_type": "minutes",
        "topic": "週次定例会議",
        "purpose": "進捗共有"
    }

    result2 = worker.execute(task2)

    if result2["status"] == "success":
        print(f"\n文書タイプ: {result2['output']['template_name']}")
        print("\nセクション構成:")
        for i, section in enumerate(result2['output']['sections'], 1):
            print(f"  {i}. {section['title']} ({section['type']})")


if __name__ == "__main__":
    main()
